"""Tool for reading various file formats (txt, pdf, docx)."""

import os
from pathlib import Path
from typing import Optional
import PyPDF2
import docx
from loguru import logger


class FileReaderTool:
    """
    Tool for reading resume and job description files in multiple formats.

    Supports .txt, .pdf, and .docx file formats with robust error handling.
    """

    SUPPORTED_EXTENSIONS = [".txt", ".pdf", ".docx", ".doc"]

    def __init__(self):
        """Initialize the file reader tool."""
        self.encoding_fallbacks = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

    def read_file(self, file_path: str) -> str:
        """
        Read a file and return its text content.

        Args:
            file_path: Path to the file to read.

        Returns:
            The text content of the file.

        Raises:
            FileNotFoundError: If the file doesn't exist.
            ValueError: If the file format is unsupported.
            Exception: For other file reading errors.
        """
        path = Path(file_path)

        # Validate file exists
        if not path.exists():
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")

        # Validate file extension
        if path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            logger.error(f"Unsupported file format: {path.suffix}")
            raise ValueError(
                f"Unsupported file format: {path.suffix}. "
                f"Supported formats: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        logger.info(f"Reading file: {file_path}")

        # Route to appropriate reader based on extension
        try:
            if path.suffix.lower() == ".pdf":
                content = self._read_pdf(path)
            elif path.suffix.lower() in [".docx", ".doc"]:
                content = self._read_docx(path)
            else:  # .txt or other text files
                content = self._read_text(path)

            if not content or not content.strip():
                logger.warning(f"File appears to be empty: {file_path}")
                raise ValueError(f"File is empty or contains no readable text: {file_path}")

            logger.success(f"Successfully read file: {file_path} ({len(content)} characters)")
            return content

        except Exception as e:
            logger.error(f"Error reading file {file_path}: {str(e)}")
            raise

    def _read_text(self, path: Path) -> str:
        """
        Read a plain text file with encoding fallback.

        Args:
            path: Path object to the text file.

        Returns:
            Content of the text file.
        """
        for encoding in self.encoding_fallbacks:
            try:
                with open(path, "r", encoding=encoding) as f:
                    content = f.read()
                logger.debug(f"Successfully read text file with encoding: {encoding}")
                return content
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Error reading text file: {str(e)}")
                raise

        raise ValueError(f"Could not read file with any supported encoding: {path}")

    def _read_pdf(self, path: Path) -> str:
        """
        Read a PDF file and extract text.

        Args:
            path: Path object to the PDF file.

        Returns:
            Extracted text from the PDF.
        """
        try:
            text_content = []

            with open(path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    try:
                        pdf_reader.decrypt("")
                    except Exception:
                        raise ValueError("PDF is password-protected and cannot be read")

                # Extract text from each page
                num_pages = len(pdf_reader.pages)
                logger.debug(f"Processing {num_pages} pages from PDF")

                for page_num in range(num_pages):
                    try:
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue

            full_text = "\n".join(text_content)

            if not full_text.strip():
                raise ValueError("No text could be extracted from the PDF")

            return full_text

        except Exception as e:
            logger.error(f"Error reading PDF file: {str(e)}")
            raise ValueError(f"Failed to read PDF file: {str(e)}")

    def _read_docx(self, path: Path) -> str:
        """
        Read a DOCX file and extract text.

        Args:
            path: Path object to the DOCX file.

        Returns:
            Extracted text from the DOCX file.
        """
        try:
            doc = docx.Document(path)
            paragraphs = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)

            full_text = "\n".join(paragraphs)

            if not full_text.strip():
                raise ValueError("No text could be extracted from the DOCX file")

            return full_text

        except Exception as e:
            logger.error(f"Error reading DOCX file: {str(e)}")
            raise ValueError(f"Failed to read DOCX file: {str(e)}")

    def validate_file(self, file_path: str) -> bool:
        """
        Validate if a file exists and has a supported extension.

        Args:
            file_path: Path to the file to validate.

        Returns:
            True if valid, False otherwise.
        """
        path = Path(file_path)
        return path.exists() and path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def get_file_info(self, file_path: str) -> dict:
        """
        Get metadata about a file.

        Args:
            file_path: Path to the file.

        Returns:
            Dictionary containing file metadata.
        """
        path = Path(file_path)

        if not path.exists():
            return {"exists": False}

        stat = path.stat()

        return {
            "exists": True,
            "name": path.name,
            "extension": path.suffix,
            "size_bytes": stat.st_size,
            "size_kb": round(stat.st_size / 1024, 2),
            "is_supported": path.suffix.lower() in self.SUPPORTED_EXTENSIONS,
        }
